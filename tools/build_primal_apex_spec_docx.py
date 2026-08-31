from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "specs" / "PRIMAL_APEX_IMPLEMENTATION_SPEC.md"
OUT_PATH = ROOT / "docs" / "specs" / "Primal_Apex_Character_Bot_Implementation_Spec.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "B2873B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)


def set_run_font(run, *, name="Calibri", east_asia="Malgun Gothic", size=None,
                 color=None, bold=None, italic=None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", east_asia="Malgun Gothic", size=11,
                   color="000000", bold=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_spacing(style, before, after, line=1.25) -> None:
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_numbering_definition(doc: Document, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_inline_markdown(paragraph, text: str, *, size=None, color=None) -> None:
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", east_asia="Malgun Gothic", size=(size or 11) - 0.5,
                         color=DARK_BLUE)
            set_cell_like_run_shading(run, "EEF2F6")
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def set_cell_like_run_shading(run, fill: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    r_pr.append(shd)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent_dxa=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Consolas", east_asia="Malgun Gothic", size=8.5, color="25313C")
        if idx < len(lines) - 1:
            run.add_break()
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def choose_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    n = len(headers)
    if n == 1:
        return [9360]
    if n == 2:
        return [2700, 6660]
    if n == 3:
        return [1900, 3660, 3800]
    if n == 4:
        # ID, name, detail, status-like column
        if headers[0].strip() == "ID":
            return [900, 2200, 4700, 1560]
        return [1500, 2500, 3600, 1760]
    if n == 5:
        return [800, 1700, 3900, 1500, 1460]
    base = 9360 // n
    return [base] * (n - 1) + [9360 - base * (n - 1)]


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    parsed = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        parsed.append(cells)
    headers = parsed[0]
    rows = parsed[2:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = choose_widths(headers, rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_inline_markdown(p, text, size=8.5, color=INK)
        for run in p.runs:
            run.bold = True

    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if col_idx in (0, len(row_data) - 1) and len(text) < 20:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_markdown(p, text, size=8.2, color="20252B")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def configure_document(doc: Document) -> tuple[int, int]:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11)
    set_paragraph_spacing(normal, 0, 6, 1.25)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        set_paragraph_spacing(style, before, after, 1.05)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    code_style = doc.styles.add_style("Code Inline", WD_STYLE_TYPE.CHARACTER)
    set_style_font(code_style, name="Consolas", east_asia="Malgun Gothic", size=9,
                   color=DARK_BLUE)

    bullet_num_id = add_numbering_definition(doc, True)
    decimal_num_id = add_numbering_definition(doc, False)
    return bullet_num_id, decimal_num_id


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_running_furniture(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("PRIMAL APEX  |  IMPLEMENTATION SPECIFICATION")
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        footer = section.footer
        fp = footer.paragraphs[0]
        add_page_field(fp)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(116)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IMPLEMENTATION SPECIFICATION")
    set_run_font(run, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Primal Apex")
    set_run_font(run, size=31, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RisuAI 캐릭터 봇 · Lua/LLM 하이브리드 아키텍처")
    set_run_font(run, size=14, color=DARK_BLUE)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2600, 6760])
    metadata = [
        ("문서 상태", "설계 초안 1.0"),
        ("기준 구현", "projects/battle-sim-lua / projects/battle-sim-poc"),
        ("기준일", "2026-08-31"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_margins(cell, top=120, bottom=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        lp = row.cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        add_inline_markdown(lp, label, size=9, color=INK)
        for run in lp.runs:
            run.bold = True
        vp = row.cells[1].paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        add_inline_markdown(vp, value, size=9, color="252A30")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Lua가 게임의 사실을 결정하고, LLM이 그 사실을 표현한다.")
    set_run_font(run, size=11, color=MUTED, italic=True)
    doc.add_page_break()


def add_static_toc(doc: Document, headings: list[str]) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("문서 구성")
    for idx, heading in enumerate(headings, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.35))
        run = p.add_run(f"{idx:02d}  {heading}")
        set_run_font(run, size=10.5, color=INK)
    doc.add_page_break()


def extract_main_headings(lines: list[str]) -> list[str]:
    result = []
    for line in lines:
        if line.startswith("## "):
            result.append(line[3:].strip())
    return result


def convert_markdown(doc: Document, lines: list[str], bullet_num_id: int,
                     decimal_num_id: int) -> None:
    i = 0
    in_code = False
    code_lines: list[str] = []
    first_title_skipped = False

    while i < len(lines):
        line = lines[i].rstrip("\n")

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("| ") and i + 1 < len(lines) and re.match(r"^\|[-: |]+\|$", lines[i + 1].strip()):
            block = [line, lines[i + 1].rstrip("\n")]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i].rstrip("\n"))
                i += 1
            add_markdown_table(doc, block)
            continue

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            if not first_title_skipped:
                first_title_skipped = True
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(p, line[3:].strip(), color=BLUE)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(p, line[4:].strip(), color=BLUE)
            i += 1
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_markdown(p, line[5:].strip(), color=DARK_BLUE)
            i += 1
            continue

        if line.startswith((
            "> 문서 상태:",
            "> 대상 플랫폼:",
            "> 기준 전투 코어:",
            "> 작성 기준일:",
        )):
            i += 1
            continue

        if line.startswith("> "):
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            set_table_geometry(table, [9360], indent_dxa=180)
            cell = table.cell(0, 0)
            set_cell_shading(cell, CALLOUT)
            set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_markdown(p, line[2:].strip(), size=10.5, color=INK)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if line.startswith(">"):
            i += 1
            continue

        if re.match(r"^- \[[ xX]\] ", line):
            checked = line[3].lower() == "x"
            text = line[6:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.24)
            p.paragraph_format.first_line_indent = Inches(-0.24)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("☒ " if checked else "☐ ")
            set_run_font(run, name="Segoe UI Symbol", east_asia="Malgun Gothic", size=10.5,
                         color=BLUE)
            add_inline_markdown(p, text)
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph()
            apply_numbering(p, bullet_num_id)
            add_inline_markdown(p, line[2:].strip())
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if number_match:
            p = doc.add_paragraph()
            apply_numbering(p, decimal_num_id)
            add_inline_markdown(p, number_match.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_markdown(p, line)
        i += 1


def main() -> None:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    bullet_num_id, decimal_num_id = configure_document(doc)
    add_cover(doc)
    add_static_toc(doc, extract_main_headings(lines))
    convert_markdown(doc, lines, bullet_num_id, decimal_num_id)
    add_running_furniture(doc)

    doc.core_properties.title = "Primal Apex 캐릭터 봇 구현 명세"
    doc.core_properties.subject = "RisuAI Lua/LLM 하이브리드 캐릭터 봇 아키텍처"
    doc.core_properties.author = "Primal Apex Project"
    doc.core_properties.keywords = "RisuAI, Lua, LLM, battle_sim, character bot"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
