from __future__ import annotations

import hashlib
import re
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "projects" / "roguelike-stage4a" / "RogueLikePOC.lua"
CHARX = ROOT / "characters" / "useful-bots" / "roguelikePOC-stage4A.charx"
OUTPUT = ROOT / "docs" / "reports" / "RogueLikePOC_Stage4A_코드_요약.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C8D1DC"
BODY_FONT = "Malgun Gothic"
MONO_FONT = "Consolas"


def compute_metrics():
    source_bytes = SOURCE.read_bytes()
    source_text = source_bytes.decode("utf-8")
    physical_lines = source_text.splitlines()
    skill_pattern = re.compile(
        r"(?m)^  (power_strike|whirlwind|united_strike|fireball|blizzard|arcane_storm) = "
    )
    return {
        "물리적 줄 수": f"{len(physical_lines):,}줄",
        "공백 제외 줄 수": f"{sum(bool(line.strip()) for line in physical_lines):,}줄",
        "문자 수": f"{len(source_text):,}자",
        "UTF-8 파일 크기": f"{len(source_bytes):,} bytes",
        "함수 정의/표현식": f"{len(re.findall(r'\bfunction\b', source_text))}개",
        "액티브 스킬 정의": f"{len(skill_pattern.findall(source_text))}개",
        "Lua 소스 SHA-256": hashlib.sha256(source_bytes).hexdigest().upper(),
        "Stage4A CHARX SHA-256": hashlib.sha256(CHARX.read_bytes()).hexdigest().upper(),
    }


def set_run_font(run, *, name=BODY_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_borders(table)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_numbering_definition(doc, *, num_fmt, level_text, left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(tabs)
    p_pr.append(ind)
    level.extend([start, fmt, text, suff, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = BODY_FONT
    title._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    title.font.size = Pt(23)
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.font.bold = True
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)

    subtitle = styles["Subtitle"]
    subtitle.font.name = BODY_FONT
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_metric_table(doc):
    metrics = compute_metrics()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "측정 항목"
    table.rows[0].cells[1].text = "기록값"
    for label, value in metrics.items():
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_table_geometry(table, [2700, 6660])
    mark_header_row(table.rows[0])

    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=10, color=DARK_BLUE, bold=True)
    for row in table.rows[1:]:
        for col_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, name=MONO_FONT if "SHA-256" in row.cells[0].text and col_index == 1 else BODY_FONT, size=8.7 if "SHA-256" in row.cells[0].text and col_index == 1 else 9.5)
    return table


def add_skill_table(doc):
    rows = [
        ("검사", "강타", "SP 20 / 단일", "자기 공격력 x1.8"),
        ("검사", "휩쓸기", "SP 35 / 전체", "자기 공격력 x0.9"),
        ("검사", "연합 강습", "SP 45 / 단일", "생존 파티 공격력 합계 x1.1"),
        ("마법사", "화염구", "MP 20 / 단일", "자기 공격력 x1.7"),
        ("마법사", "눈보라", "MP 35 / 전체", "자기 공격력 x0.85"),
        ("마법사", "비전 폭풍", "MP 55 / 전체", "생존 파티 공격력 합계 x0.75"),
    ]
    table = doc.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ("소유자", "스킬", "비용 / 대상", "공격 스케일링")):
        cell.text = text
    for row_data in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row_data):
            cells[index].text = text
    set_table_geometry(table, [1050, 1550, 2050, 4710])
    mark_header_row(table.rows[0])

    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    for row in table.rows[1:]:
        for col_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index < 3 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
    return table


def build_document():
    if not SOURCE.exists() or not CHARX.exists():
        raise FileNotFoundError("Stage4A source or CHARX deliverable is missing.")

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "ROGUELIKEPOC  /  STAGE 4-A CODE BRIEF"
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)
    add_page_field(section.footer.paragraphs[0])

    bullet_num_id = add_numbering_definition(doc, num_fmt="bullet", level_text="•")
    decimal_num_id = add_numbering_definition(doc, num_fmt="decimal", level_text="%1.")

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("TECHNICAL REFERENCE")
    set_run_font(run, size=9, color=BLUE, bold=True)

    doc.add_paragraph("RogueLikePOC Stage 4-A 코드 요약", style="Title")
    doc.add_paragraph("SP/MP 및 공격형 액티브 스킬 전투 엔진", style="Subtitle")

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(12)
    metadata_run = metadata.add_run(f"작성일 {date.today().isoformat()}  |  기준 소스 RogueLikePOC.lua  |  결과물 roguelikePOC-stage4A.charx")
    set_run_font(metadata_run, size=9, color=MUTED)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(8)
    run = lead.add_run("요약. ")
    set_run_font(run, size=11, color=DARK_BLUE, bold=True)
    run = lead.add_run("기존 기본 공격 POC를 데이터 기반 스킬 엔진으로 확장했다. 이번 범위는 SP/MP 자원과 4-A형 단일·전체 공격 스킬이며, 버프·디버프는 다음 단계로 남겨 두었다.")
    set_run_font(run, size=11)

    doc.add_paragraph("코드 메트릭", style="Heading 1")
    add_metric_table(doc)
    note = doc.add_paragraph("측정 기준: 물리적 줄 수는 파일의 실제 라인 수, 문자 수는 UTF-8 디코딩 후 Unicode 문자 수, 파일 크기는 디스크상의 UTF-8 바이트 수입니다.")
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(4)
    for run in note.runs:
        set_run_font(run, size=8.5, color=MUTED, italic=True)

    doc.add_paragraph("코드 구성", style="Heading 1")
    architecture = [
        "ACTIVE_SKILLS 테이블: 비용 자원, 비용, 대상 범위, 공격력 기준, 배율을 선언하는 데이터 계층",
        "상태 정규화: 이전 POC의 세이브를 상태 버전 2로 보정하고 SP/MP·스킬 목록을 채움",
        "맵 계층: 1-2-3-2-1 레이어 DAG와 전투·이벤트·상점·보스 방 생성",
        "전투 계층: 기본 공격, 단일/전체 스킬, 자원 차감, 승패 판정, 파티 상태 동기화",
        "렌더링 계층: 파티 자원, 전투 유닛, 스킬 메뉴, 대상 선택, 전투 로그를 HTML로 생성",
        "RisuAI 콜백: editDisplay, onStart, onButtonClick을 통해 화면과 상태 전이를 연결",
    ]
    for item in architecture:
        add_numbered_paragraph(doc, item, bullet_num_id)

    schema = doc.add_paragraph()
    schema.paragraph_format.space_before = Pt(6)
    schema.paragraph_format.space_after = Pt(8)
    schema.paragraph_format.left_indent = Inches(0.12)
    schema.paragraph_format.right_indent = Inches(0.12)
    set_paragraph_shading(schema, LIGHT_GRAY)
    run = schema.add_run("skill = { costType, cost, target, scaling, multiplier }\n")
    set_run_font(run, name=MONO_FONT, size=9.5, color=DARK_BLUE, bold=True)
    run = schema.add_run("damage = max(0, round(baseAttack x multiplier) - targetDefense)")
    set_run_font(run, name=MONO_FONT, size=9.5)

    doc.add_page_break()

    doc.add_paragraph("액티브 스킬 카탈로그", style="Heading 1")
    add_skill_table(doc)

    doc.add_paragraph("플레이어 행동 흐름", style="Heading 1")
    flow = [
        "행동할 생존 파티원을 선택한다.",
        "기본 공격 또는 보유 액티브 스킬 중 하나를 선택한다.",
        "단일 대상 행동은 적을 지정하고, 전체 대상 스킬은 즉시 실행한다.",
        "스킬 비용을 한 번 차감하고 각 대상에게 배율 공격과 방어력 계산을 적용한다.",
        "적이 생존했다면 기존 랜덤 타게팅 AI가 행동하고 다음 라운드로 넘어간다.",
    ]
    for item in flow:
        add_numbered_paragraph(doc, item, decimal_num_id)

    doc.add_paragraph("구현 범위", style="Heading 1")
    implemented = [
        "SP/MP와 최대치, 전투 간 자원 유지, 자원 부족 스킬 비활성화",
        "스킬 개수 제한 없는 ID 배열, 단일/전체 대상, 자기/파티 합산 공격력 스케일링",
        "선택 취소, 대상 유효성 검사, 전투 로그, 이전 상태 마이그레이션",
    ]
    for item in implemented:
        add_numbered_paragraph(doc, item, bullet_num_id)

    deferred = doc.add_paragraph()
    deferred.paragraph_format.space_before = Pt(6)
    deferred.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(deferred, LIGHT_GRAY)
    run = deferred.add_run("다음 단계: ")
    set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    run = deferred.add_run("공격/방어 감소, 침묵, 가속, 중독, 도발과 지속 턴 처리. SP/MP 회복 규칙도 아직 정의하지 않았다.")
    set_run_font(run, size=10.5)

    doc.add_paragraph("빌드 및 검증", style="Heading 1")
    verification = [
        "원본 CHARX와 별도 백업의 SHA-256 일치 확인",
        "module.risum 재디코딩 및 risuModule 구조 확인",
        "임베디드 Lua와 작업 소스의 SHA-256 완전 일치",
        "Lua 문자열·괄호·블록 구조 검증 통과",
        "최종 런타임 검증은 RisuAI에 CHARX를 임포트해 수행 필요",
    ]
    for item in verification:
        add_numbered_paragraph(doc, item, bullet_num_id)

    source_path = doc.add_paragraph()
    source_path.paragraph_format.space_before = Pt(6)
    source_path.paragraph_format.space_after = Pt(0)
    run = source_path.add_run(f"Source: {SOURCE}")
    set_run_font(run, name=MONO_FONT, size=8.5, color=MUTED)
    output_path = doc.add_paragraph()
    output_path.paragraph_format.space_before = Pt(0)
    output_path.paragraph_format.space_after = Pt(0)
    run = output_path.add_run(f"CHARX:  {CHARX}")
    set_run_font(run, name=MONO_FONT, size=8.5, color=MUTED)

    doc.core_properties.title = "RogueLikePOC Stage 4-A 코드 요약"
    doc.core_properties.subject = "SP/MP 및 공격형 액티브 스킬 전투 엔진 기술 요약"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "RisuAI, RogueLikePOC, Lua, Stage4A, active skills"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
